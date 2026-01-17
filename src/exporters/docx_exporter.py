"""
Word Document Exporter Module

Exports clinical documents as Word (.docx) files with optional
letterhead support.
"""

from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
from utils.structured_logging import get_logger

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from exporters.base_exporter import BaseExporter

logger = get_logger(__name__)


class DocxExporter(BaseExporter):
    """Word document exporter.

    Exports clinical documents as .docx files with support for
    SOAP notes, referrals, letters, and other document types.
    """

    def __init__(
        self,
        clinic_name: str = "",
        doctor_name: str = ""
    ):
        """Initialize the Word exporter.

        Args:
            clinic_name: Name of the clinic for letterhead
            doctor_name: Name of the doctor for letterhead
        """
        super().__init__()
        self.clinic_name = clinic_name
        self.doctor_name = doctor_name

    def set_letterhead(self, clinic_name: str, doctor_name: str) -> None:
        """Set letterhead information.

        Args:
            clinic_name: Name of the clinic
            doctor_name: Name of the doctor
        """
        self.clinic_name = clinic_name
        self.doctor_name = doctor_name

    def export(
        self,
        content: Dict[str, Any],
        output_path: Path
    ) -> bool:
        """Export content to a Word document.

        Args:
            content: Dictionary containing the document content.
                Expected keys:
                - document_type: Type of document (soap, referral, letter, generic)
                - content: Text content (string or dict with sections)
                - title: Optional document title
                - include_letterhead: Whether to include letterhead
                - patient_info: Optional patient information

        Returns:
            True if export was successful, False otherwise.
        """
        try:
            # Ensure directory exists
            if not self._ensure_directory(output_path):
                return False

            # Create document
            document = Document()

            # Set up styles
            self._setup_styles(document)

            # Add letterhead if requested
            if content.get("include_letterhead", False):
                self._add_letterhead(document)

            # Get document type
            doc_type = content.get("document_type", "generic")

            # Route to appropriate export method
            if doc_type == "soap":
                self._add_soap_content(document, content)
            elif doc_type == "referral":
                self._add_referral_content(document, content)
            elif doc_type == "letter":
                self._add_letter_content(document, content)
            else:
                self._add_generic_content(document, content)

            # Save document
            document.save(str(output_path))
            logger.info(f"Word document exported to: {output_path}")
            return True

        except Exception as e:
            self._last_error = f"Failed to export Word document: {str(e)}"
            logger.error(self._last_error, exc_info=True)
            return False

    def export_to_string(self, content: Dict[str, Any]) -> str:
        """Export content as string representation.

        Word documents can't be meaningfully represented as a string,
        so this returns a text representation of the content.

        Args:
            content: Document content dictionary

        Returns:
            Text representation of the content
        """
        text_content = content.get("content", "")
        if isinstance(text_content, dict):
            # Combine sections
            sections = []
            for section in ["subjective", "objective", "assessment", "plan"]:
                if section in text_content and text_content[section]:
                    sections.append(f"{section.upper()}:\n{text_content[section]}")
            return "\n\n".join(sections)
        return str(text_content)

    def _setup_styles(self, document: Document) -> None:
        """Set up document styles.

        Args:
            document: Word document to configure
        """
        styles = document.styles

        # Title style
        try:
            title_style = styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = "Arial"
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.paragraph_format.space_after = Pt(12)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except ValueError:
            pass  # Style already exists

        # Section header style
        try:
            header_style = styles.add_style("SectionHeader", WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = "Arial"
            header_style.font.size = Pt(12)
            header_style.font.bold = True
            header_style.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
            header_style.paragraph_format.space_before = Pt(12)
            header_style.paragraph_format.space_after = Pt(6)
        except ValueError:
            pass  # Style already exists

        # Normal text style
        try:
            normal_style = styles.add_style("ContentText", WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = "Arial"
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.15
        except ValueError:
            pass  # Style already exists

    def _add_letterhead(self, document: Document) -> None:
        """Add simple letterhead with clinic and doctor name.

        Args:
            document: Word document to add letterhead to
        """
        if not self.clinic_name and not self.doctor_name:
            return

        # Add clinic name
        if self.clinic_name:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(self.clinic_name)
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True

        # Add doctor name
        if self.doctor_name:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(self.doctor_name)
            run.font.name = "Arial"
            run.font.size = Pt(11)

        # Add separator line
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_" * 60)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Add some space
        document.add_paragraph()

    def _add_soap_content(self, document: Document, content: Dict[str, Any]) -> None:
        """Add SOAP note content to document.

        Args:
            document: Word document
            content: Content dictionary
        """
        # Add title
        title = content.get("title", "SOAP Note")
        p = document.add_paragraph(title)
        p.style = "Title"

        # Add date
        date_str = datetime.now().strftime("%B %d, %Y")
        p = document.add_paragraph(f"Date: {date_str}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add patient info if provided
        patient_info = content.get("patient_info", {})
        if patient_info:
            self._add_patient_info(document, patient_info)

        # Add SOAP sections
        soap_content = content.get("content", {})

        if isinstance(soap_content, str):
            # Parse SOAP sections from text
            soap_content = self._parse_soap_text(soap_content)

        section_titles = {
            "subjective": "SUBJECTIVE",
            "objective": "OBJECTIVE",
            "assessment": "ASSESSMENT",
            "plan": "PLAN"
        }

        for section_key, section_title in section_titles.items():
            section_text = soap_content.get(section_key, "")
            if section_text:
                # Add section header
                p = document.add_paragraph(section_title)
                try:
                    p.style = "SectionHeader"
                except KeyError:
                    p.runs[0].bold = True

                # Add section content
                for para in section_text.split("\n"):
                    if para.strip():
                        p = document.add_paragraph(para)
                        try:
                            p.style = "ContentText"
                        except KeyError:
                            pass

    def _add_referral_content(self, document: Document, content: Dict[str, Any]) -> None:
        """Add referral letter content to document.

        Args:
            document: Word document
            content: Content dictionary
        """
        # Add title
        title = content.get("title", "Referral Letter")
        p = document.add_paragraph(title)
        p.style = "Title"

        # Add date
        date_str = datetime.now().strftime("%B %d, %Y")
        p = document.add_paragraph(date_str)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_paragraph()  # Space

        # Add content
        text_content = content.get("content", "")
        if isinstance(text_content, dict):
            text_content = text_content.get("content", "")

        for para in text_content.split("\n\n"):
            if para.strip():
                p = document.add_paragraph(para.strip())
                try:
                    p.style = "ContentText"
                except KeyError:
                    pass

    def _add_letter_content(self, document: Document, content: Dict[str, Any]) -> None:
        """Add medical letter content to document.

        Args:
            document: Word document
            content: Content dictionary
        """
        # Add date
        date_str = datetime.now().strftime("%B %d, %Y")
        p = document.add_paragraph(date_str)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_paragraph()  # Space

        # Add content
        text_content = content.get("content", "")
        if isinstance(text_content, dict):
            text_content = text_content.get("content", "")

        for para in text_content.split("\n\n"):
            if para.strip():
                p = document.add_paragraph(para.strip())
                try:
                    p.style = "ContentText"
                except KeyError:
                    pass

    def _add_generic_content(self, document: Document, content: Dict[str, Any]) -> None:
        """Add generic document content.

        Args:
            document: Word document
            content: Content dictionary
        """
        # Add title if provided
        title = content.get("title")
        if title:
            p = document.add_paragraph(title)
            p.style = "Title"

        # Add date
        date_str = datetime.now().strftime("%B %d, %Y")
        p = document.add_paragraph(date_str)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_paragraph()  # Space

        # Add content
        text_content = content.get("content", "")
        if isinstance(text_content, dict):
            # Handle structured content
            for key, value in text_content.items():
                if value:
                    p = document.add_paragraph(f"{key.title()}:")
                    p.runs[0].bold = True
                    p = document.add_paragraph(str(value))
        else:
            # Plain text content
            for para in str(text_content).split("\n\n"):
                if para.strip():
                    p = document.add_paragraph(para.strip())

    def _add_patient_info(self, document: Document, patient_info: Dict[str, Any]) -> None:
        """Add patient information section.

        Args:
            document: Word document
            patient_info: Patient information dictionary
        """
        if not patient_info:
            return

        # Add patient info
        info_parts = []
        if patient_info.get("name"):
            info_parts.append(f"Patient: {patient_info['name']}")
        if patient_info.get("dob"):
            info_parts.append(f"DOB: {patient_info['dob']}")
        if patient_info.get("id"):
            info_parts.append(f"ID: {patient_info['id']}")

        if info_parts:
            p = document.add_paragraph(" | ".join(info_parts))
            p.runs[0].font.size = Pt(10)
            document.add_paragraph()  # Space

    def _parse_soap_text(self, text: str) -> Dict[str, str]:
        """Parse SOAP note text into sections.

        Args:
            text: Full SOAP note text

        Returns:
            Dictionary with section keys
        """
        sections = {
            "subjective": "",
            "objective": "",
            "assessment": "",
            "plan": ""
        }

        current_section = None
        current_content = []

        section_patterns = {
            "subjective": ["subjective", "s:", "s.", "chief complaint", "cc:"],
            "objective": ["objective", "o:", "o.", "physical exam", "pe:"],
            "assessment": ["assessment", "a:", "a.", "impression", "diagnoses"],
            "plan": ["plan", "p:", "p.", "treatment", "recommendations"]
        }

        for line in text.split("\n"):
            line_lower = line.lower().strip()

            # Check if this line is a section header
            found_section = None
            for section_name, patterns in section_patterns.items():
                for pattern in patterns:
                    if line_lower == pattern or line_lower.startswith(pattern + ":") or line_lower.startswith(pattern + " "):
                        found_section = section_name
                        break
                if found_section:
                    break

            if found_section:
                # Save previous section
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = found_section
                current_content = []
            elif current_section:
                current_content.append(line)

        # Save last section
        if current_section and current_content:
            sections[current_section] = "\n".join(current_content).strip()

        # If no sections found, put everything in subjective
        if not any(sections.values()):
            sections["subjective"] = text

        return sections

    def export_soap_note(
        self,
        soap_text: str,
        output_path: Path,
        title: str = "SOAP Note",
        include_letterhead: bool = False,
        patient_info: Optional[Dict[str, Any]] = None
    ) -> bool:
        """Export a SOAP note as Word document.

        Convenience method for exporting SOAP notes.

        Args:
            soap_text: Full SOAP note text
            output_path: Path to save the document
            title: Document title
            include_letterhead: Whether to include letterhead
            patient_info: Optional patient information

        Returns:
            True if successful, False otherwise
        """
        content = {
            "document_type": "soap",
            "content": soap_text,
            "title": title,
            "include_letterhead": include_letterhead,
            "patient_info": patient_info
        }
        return self.export(content, output_path)

    def export_referral(
        self,
        referral_text: str,
        output_path: Path,
        title: str = "Referral Letter",
        include_letterhead: bool = False
    ) -> bool:
        """Export a referral letter as Word document.

        Args:
            referral_text: Referral letter text
            output_path: Path to save the document
            title: Document title
            include_letterhead: Whether to include letterhead

        Returns:
            True if successful, False otherwise
        """
        content = {
            "document_type": "referral",
            "content": referral_text,
            "title": title,
            "include_letterhead": include_letterhead
        }
        return self.export(content, output_path)

    def export_letter(
        self,
        letter_text: str,
        output_path: Path,
        title: str = "Medical Letter",
        include_letterhead: bool = False
    ) -> bool:
        """Export a medical letter as Word document.

        Args:
            letter_text: Letter text
            output_path: Path to save the document
            title: Document title
            include_letterhead: Whether to include letterhead

        Returns:
            True if successful, False otherwise
        """
        content = {
            "document_type": "letter",
            "content": letter_text,
            "title": title,
            "include_letterhead": include_letterhead
        }
        return self.export(content, output_path)


def get_docx_exporter(
    clinic_name: str = "",
    doctor_name: str = ""
) -> DocxExporter:
    """Get a Word document exporter instance.

    Factory function for creating Word exporter with optional letterhead.

    Args:
        clinic_name: Name of the clinic for letterhead
        doctor_name: Name of the doctor for letterhead

    Returns:
        DocxExporter instance
    """
    return DocxExporter(clinic_name=clinic_name, doctor_name=doctor_name)
