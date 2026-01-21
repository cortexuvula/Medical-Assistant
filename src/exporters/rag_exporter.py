"""
RAG Conversation Exporter.

Exports RAG conversations to PDF, Word, Markdown, and JSON formats.
Includes query/response pairs, source citations, and metadata.
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from exporters.base_exporter import BaseExporter
from utils.structured_logging import get_logger

logger = get_logger(__name__)


class RAGConversationExporter(BaseExporter):
    """Exports RAG conversations to multiple formats."""

    def __init__(self):
        """Initialize the RAG conversation exporter."""
        super().__init__()
        self._pdf_available = False
        self._docx_available = False

        # Check for optional dependencies
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            self._pdf_available = True
        except ImportError:
            logger.debug("ReportLab not available, PDF export disabled")

        try:
            from docx import Document
            self._docx_available = True
        except ImportError:
            logger.debug("python-docx not available, Word export disabled")

    def export(self, content: Dict[str, Any], output_path: Path) -> bool:
        """Export conversation based on file extension.

        Args:
            content: Conversation data
            output_path: Output file path

        Returns:
            True if export succeeded
        """
        if not self._ensure_directory(output_path):
            return False

        ext = output_path.suffix.lower()

        if ext == '.pdf':
            return self.export_as_pdf(content, output_path)
        elif ext == '.docx':
            return self.export_as_docx(content, output_path)
        elif ext == '.md':
            return self.export_as_markdown(content, output_path)
        elif ext == '.json':
            return self.export_as_json(content, output_path)
        else:
            self._last_error = f"Unsupported export format: {ext}"
            return False

    def export_to_string(self, content: Dict[str, Any]) -> str:
        """Export conversation as markdown string.

        Args:
            content: Conversation data

        Returns:
            Markdown-formatted string
        """
        return self._format_as_markdown(content)

    def export_as_pdf(self, content: Dict[str, Any], output_path: Path) -> bool:
        """Export conversation as PDF document.

        Args:
            content: Conversation data
            output_path: Output PDF path

        Returns:
            True if export succeeded
        """
        if not self._pdf_available:
            self._last_error = "ReportLab not available. Install with: pip install reportlab"
            return False

        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                Table,
                TableStyle,
            )

            # Create document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=12,
                spaceBefore=12,
                spaceAfter=6,
                textColor=colors.HexColor('#2196F3')
            )
            query_style = ParagraphStyle(
                'Query',
                parent=styles['Normal'],
                fontSize=11,
                spaceBefore=6,
                spaceAfter=6,
                leftIndent=0,
                textColor=colors.HexColor('#333333'),
                backColor=colors.HexColor('#E3F2FD')
            )
            response_style = ParagraphStyle(
                'Response',
                parent=styles['Normal'],
                fontSize=10,
                spaceBefore=6,
                spaceAfter=12,
                leftIndent=20
            )
            source_style = ParagraphStyle(
                'Source',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.gray,
                leftIndent=40
            )

            # Build content
            story = []

            # Title
            title = content.get('title', 'RAG Conversation Export')
            story.append(Paragraph(title, title_style))

            # Metadata
            timestamp = content.get('timestamp', datetime.now().isoformat())
            story.append(Paragraph(f"Exported: {timestamp}", styles['Normal']))
            story.append(Spacer(1, 20))

            # Summary
            metadata = content.get('metadata', {})
            if metadata:
                summary_text = f"Total queries: {metadata.get('total_queries', 0)}"
                if metadata.get('documents_searched'):
                    summary_text += f" | Documents searched: {metadata.get('documents_searched')}"
                story.append(Paragraph(summary_text, styles['Normal']))
                story.append(Spacer(1, 20))

            # Exchanges
            exchanges = content.get('exchanges', [])
            for i, exchange in enumerate(exchanges, 1):
                # Query header
                story.append(Paragraph(f"Query {i}", heading_style))

                # Query text
                query = exchange.get('query', '')
                story.append(Paragraph(f"<b>Q:</b> {self._escape_html(query)}", query_style))

                # Response
                response = exchange.get('response', '')
                story.append(Paragraph(self._escape_html(response), response_style))

                # Sources
                sources = exchange.get('sources', [])
                if sources:
                    story.append(Paragraph("<b>Sources:</b>", source_style))
                    for source in sources:
                        source_text = f"- {source.get('document', 'Unknown')}"
                        if source.get('score'):
                            source_text += f" ({source['score']:.1%})"
                        story.append(Paragraph(source_text, source_style))

                story.append(Spacer(1, 15))

            # Build PDF
            doc.build(story)
            logger.info(f"RAG conversation exported to PDF: {output_path}")
            return True

        except Exception as e:
            self._last_error = f"PDF export failed: {str(e)}"
            logger.error(self._last_error)
            return False

    def export_as_docx(self, content: Dict[str, Any], output_path: Path) -> bool:
        """Export conversation as Word document.

        Args:
            content: Conversation data
            output_path: Output DOCX path

        Returns:
            True if export succeeded
        """
        if not self._docx_available:
            self._last_error = "python-docx not available. Install with: pip install python-docx"
            return False

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Create document
            doc = Document()

            # Title
            title = content.get('title', 'RAG Conversation Export')
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            timestamp = content.get('timestamp', datetime.now().isoformat())
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run(f"Exported: {timestamp}")
            meta_run.italic = True
            meta_run.font.color.rgb = RGBColor(128, 128, 128)

            metadata = content.get('metadata', {})
            if metadata:
                meta_para2 = doc.add_paragraph()
                summary = f"Total queries: {metadata.get('total_queries', 0)}"
                if metadata.get('documents_searched'):
                    summary += f" | Documents searched: {metadata.get('documents_searched')}"
                meta_para2.add_run(summary)

            doc.add_paragraph()  # Spacer

            # Exchanges
            exchanges = content.get('exchanges', [])
            for i, exchange in enumerate(exchanges, 1):
                # Query heading
                doc.add_heading(f"Query {i}", level=1)

                # Query paragraph
                query_para = doc.add_paragraph()
                query_run = query_para.add_run("Q: ")
                query_run.bold = True
                query_para.add_run(exchange.get('query', ''))

                # Response paragraph
                response = exchange.get('response', '')
                response_para = doc.add_paragraph()
                response_para.add_run(response)
                response_para.paragraph_format.left_indent = Inches(0.25)

                # Sources
                sources = exchange.get('sources', [])
                if sources:
                    sources_heading = doc.add_paragraph()
                    sources_run = sources_heading.add_run("Sources:")
                    sources_run.bold = True
                    sources_run.font.size = Pt(10)

                    for source in sources:
                        source_para = doc.add_paragraph(style='List Bullet')
                        source_text = source.get('document', 'Unknown')
                        if source.get('score'):
                            source_text += f" ({source['score']:.1%})"
                        source_run = source_para.add_run(source_text)
                        source_run.font.size = Pt(9)
                        source_run.font.color.rgb = RGBColor(128, 128, 128)

                doc.add_paragraph()  # Spacer between exchanges

            # Save document
            doc.save(str(output_path))
            logger.info(f"RAG conversation exported to Word: {output_path}")
            return True

        except Exception as e:
            self._last_error = f"Word export failed: {str(e)}"
            logger.error(self._last_error)
            return False

    def export_as_markdown(self, content: Dict[str, Any], output_path: Path) -> bool:
        """Export conversation as Markdown file.

        Args:
            content: Conversation data
            output_path: Output MD path

        Returns:
            True if export succeeded
        """
        try:
            markdown = self._format_as_markdown(content)
            output_path.write_text(markdown, encoding='utf-8')
            logger.info(f"RAG conversation exported to Markdown: {output_path}")
            return True
        except Exception as e:
            self._last_error = f"Markdown export failed: {str(e)}"
            logger.error(self._last_error)
            return False

    def export_as_json(self, content: Dict[str, Any], output_path: Path) -> bool:
        """Export conversation as JSON file.

        Args:
            content: Conversation data
            output_path: Output JSON path

        Returns:
            True if export succeeded
        """
        try:
            # Ensure datetime objects are serializable
            def json_serializer(obj):
                if isinstance(obj, datetime):
                    return obj.isoformat()
                raise TypeError(f"Object of type {type(obj)} is not JSON serializable")

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(content, f, indent=2, default=json_serializer)

            logger.info(f"RAG conversation exported to JSON: {output_path}")
            return True

        except Exception as e:
            self._last_error = f"JSON export failed: {str(e)}"
            logger.error(self._last_error)
            return False

    def _format_as_markdown(self, content: Dict[str, Any]) -> str:
        """Format conversation as Markdown.

        Args:
            content: Conversation data

        Returns:
            Markdown string
        """
        lines = []

        # Title
        title = content.get('title', 'RAG Conversation Export')
        lines.append(f"# {title}")
        lines.append("")

        # Metadata
        timestamp = content.get('timestamp', datetime.now().isoformat())
        lines.append(f"*Exported: {timestamp}*")
        lines.append("")

        metadata = content.get('metadata', {})
        if metadata:
            lines.append("## Summary")
            lines.append(f"- **Total queries:** {metadata.get('total_queries', 0)}")
            if metadata.get('documents_searched'):
                lines.append(f"- **Documents searched:** {metadata.get('documents_searched')}")
            if metadata.get('search_settings'):
                lines.append(f"- **Settings:** {metadata.get('search_settings')}")
            lines.append("")

        # Exchanges
        lines.append("---")
        lines.append("")

        exchanges = content.get('exchanges', [])
        for i, exchange in enumerate(exchanges, 1):
            lines.append(f"## Query {i}")
            lines.append("")

            # Query
            query = exchange.get('query', '')
            lines.append(f"**Q:** {query}")
            lines.append("")

            # Response
            response = exchange.get('response', '')
            lines.append(response)
            lines.append("")

            # Sources
            sources = exchange.get('sources', [])
            if sources:
                lines.append("### Sources")
                for source in sources:
                    source_text = f"- {source.get('document', 'Unknown')}"
                    if source.get('chunk'):
                        preview = source['chunk'][:100] + "..." if len(source.get('chunk', '')) > 100 else source.get('chunk', '')
                        source_text += f": *{preview}*"
                    if source.get('score'):
                        source_text += f" ({source['score']:.1%})"
                    lines.append(source_text)
                lines.append("")

            # Query expansion info
            expansion = exchange.get('query_expansion')
            if expansion:
                lines.append("### Query Expansion")
                lines.append(f"- **Original:** {expansion.get('original_query', query)}")
                if expansion.get('expanded_terms'):
                    lines.append(f"- **Expanded terms:** {', '.join(expansion['expanded_terms'])}")
                lines.append("")

            # Processing time
            if exchange.get('processing_time_ms'):
                lines.append(f"*Processing time: {exchange['processing_time_ms']:.0f}ms*")
                lines.append("")

            lines.append("---")
            lines.append("")

        return "\n".join(lines)

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters for PDF.

        Args:
            text: Text to escape

        Returns:
            Escaped text
        """
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))


def create_rag_export_content(
    session_id: str,
    exchanges: list,
    title: str = "RAG Conversation",
    metadata: Optional[Dict] = None
) -> Dict[str, Any]:
    """Create export content dictionary from RAG session data.

    Args:
        session_id: Session identifier
        exchanges: List of query/response exchanges
        title: Export title
        metadata: Optional additional metadata

    Returns:
        Content dictionary suitable for export
    """
    export_content = {
        "session_id": session_id,
        "title": title,
        "timestamp": datetime.now().isoformat(),
        "exchanges": [],
        "metadata": metadata or {}
    }

    for exchange in exchanges:
        exchange_data = {
            "query": exchange.get('query', ''),
            "response": exchange.get('response', ''),
            "sources": [],
            "query_expansion": exchange.get('query_expansion'),
            "processing_time_ms": exchange.get('processing_time_ms', 0)
        }

        # Add sources
        for source in exchange.get('sources', []):
            source_data = {
                "document": source.get('document_filename', source.get('document', 'Unknown')),
                "chunk": source.get('chunk_text', source.get('chunk', '')),
                "score": source.get('combined_score', source.get('score', 0))
            }
            exchange_data["sources"].append(source_data)

        export_content["exchanges"].append(exchange_data)

    # Update metadata
    export_content["metadata"]["total_queries"] = len(exchanges)

    return export_content


# Singleton instance
_exporter: Optional[RAGConversationExporter] = None


def get_rag_exporter() -> RAGConversationExporter:
    """Get the global RAG exporter instance.

    Returns:
        RAGConversationExporter instance
    """
    global _exporter
    if _exporter is None:
        _exporter = RAGConversationExporter()
    return _exporter
