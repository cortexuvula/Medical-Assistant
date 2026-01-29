"""
Document processor for RAG system.

Handles:
- Text extraction from PDF, DOCX, TXT files
- OCR for scanned documents and images (Azure Document Intelligence)
- Semantic chunking with token counting

OCR Provider:
    Azure Document Intelligence (prebuilt-read model)
"""

import hashlib
import io
from utils.structured_logging import get_logger
import os
import re
from pathlib import Path
from typing import Optional

import tiktoken

from rag.models import (
    DocumentChunk,
    DocumentMetadata,
    DocumentType,
    RAGDocument,
    UploadStatus,
)

logger = get_logger(__name__)

# Lazy import for OCR providers to avoid circular imports
_ocr_manager = None

# Supported file extensions mapped to document types
EXTENSION_TO_TYPE = {
    ".pdf": DocumentType.PDF,
    ".docx": DocumentType.DOCX,
    ".doc": DocumentType.DOCX,
    ".txt": DocumentType.TXT,
    ".md": DocumentType.TXT,
    ".png": DocumentType.IMAGE,
    ".jpg": DocumentType.IMAGE,
    ".jpeg": DocumentType.IMAGE,
    ".tiff": DocumentType.IMAGE,
    ".tif": DocumentType.IMAGE,
    ".bmp": DocumentType.IMAGE,
}


class DocumentProcessor:
    """Processes documents for RAG ingestion."""

    def __init__(
        self,
        chunk_size_tokens: int = 500,
        chunk_overlap_tokens: int = 50,
        encoding_name: str = "cl100k_base",
        ocr_language: str = "eng",
    ):
        """Initialize the document processor.

        Args:
            chunk_size_tokens: Target size for each chunk in tokens
            chunk_overlap_tokens: Overlap between chunks in tokens
            encoding_name: Tiktoken encoding name for token counting
            ocr_language: Language code for OCR (e.g., 'eng', 'fra')
        """
        self.chunk_size_tokens = chunk_size_tokens
        self.chunk_overlap_tokens = chunk_overlap_tokens
        self.ocr_language = ocr_language

        # Initialize tiktoken encoder
        try:
            self.encoder = tiktoken.get_encoding(encoding_name)
        except Exception as e:
            logger.warning(f"Failed to load tiktoken encoding {encoding_name}: {e}")
            self.encoder = None

    def count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken.

        Falls back to a character-based estimate (~4 chars per token) when
        tiktoken encodings are unavailable (e.g. PyInstaller bundle).

        Args:
            text: Text to count tokens for

        Returns:
            Number of tokens
        """
        if not text:
            return 0
        if self.encoder is not None:
            return len(self.encoder.encode(text))
        # Approximate: ~4 characters per token for English text
        return max(1, len(text) // 4)

    def _get_ocr_manager(self):
        """Get or create the OCR provider manager.

        Returns:
            OCRProviderManager instance
        """
        global _ocr_manager
        if _ocr_manager is None:
            try:
                from ocr_providers import get_ocr_manager
                _ocr_manager = get_ocr_manager()
            except ImportError as e:
                logger.warning(f"OCR providers not available: {e}")
                return None
        return _ocr_manager

    def get_document_type(self, file_path: str) -> Optional[DocumentType]:
        """Determine document type from file extension.

        Args:
            file_path: Path to the document

        Returns:
            DocumentType or None if unsupported
        """
        ext = Path(file_path).suffix.lower()
        return EXTENSION_TO_TYPE.get(ext)


    def extract_text(self, file_path: str, enable_ocr: bool = True) -> tuple[str, DocumentMetadata, int, bool]:
        """Extract text from a document.

        Args:
            file_path: Path to the document
            enable_ocr: Whether to use OCR for scanned documents/images

        Returns:
            Tuple of (extracted_text, metadata, page_count, ocr_was_used)
        """
        doc_type = self.get_document_type(file_path)
        if not doc_type:
            raise ValueError(f"Unsupported file type: {file_path}")

        if doc_type == DocumentType.PDF:
            return self._extract_from_pdf(file_path, enable_ocr)
        elif doc_type == DocumentType.DOCX:
            return self._extract_from_docx(file_path)
        elif doc_type == DocumentType.TXT:
            return self._extract_from_txt(file_path)
        elif doc_type == DocumentType.IMAGE:
            return self._extract_from_image(file_path, enable_ocr)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

    def _extract_from_pdf(self, file_path: str, enable_ocr: bool) -> tuple[str, DocumentMetadata, int, bool]:
        """Extract text from PDF using pdfplumber, with OCR for scanned pages.

        Args:
            file_path: Path to PDF file
            enable_ocr: Whether to use OCR for scanned pages

        Returns:
            Tuple of (text, metadata, page_count, ocr_used)
        """
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required for PDF processing. Install with: pip install pdfplumber")

        text_parts = []
        ocr_used = False
        metadata = DocumentMetadata()

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)

            # Extract metadata
            if pdf.metadata:
                metadata.title = pdf.metadata.get("Title")
                metadata.author = pdf.metadata.get("Author")
                metadata.subject = pdf.metadata.get("Subject")
                if pdf.metadata.get("Keywords"):
                    metadata.keywords = [k.strip() for k in pdf.metadata.get("Keywords", "").split(",") if k.strip()]

            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""

                # Check if page might be scanned (very little text extracted)
                if enable_ocr and len(page_text.strip()) < 50:
                    # Try OCR
                    ocr_text = self._ocr_pdf_page(page)
                    if ocr_text and len(ocr_text) > len(page_text):
                        page_text = ocr_text
                        ocr_used = True

                if page_text:
                    text_parts.append(f"[Page {i + 1}]\n{page_text}")

        return "\n\n".join(text_parts), metadata, page_count, ocr_used

    def _ocr_pdf_page(self, page) -> str:
        """Perform OCR on a PDF page using Azure Document Intelligence.

        Args:
            page: pdfplumber page object

        Returns:
            OCR extracted text
        """
        try:
            from PIL import Image
        except ImportError:
            logger.warning("PIL not available for PDF page rendering")
            return ""

        try:
            # Convert page to image
            img = page.to_image(resolution=300)
            pil_image = img.original

            # Use OCR provider manager
            ocr_manager = self._get_ocr_manager()
            if ocr_manager:
                result = ocr_manager.extract_from_pil_image(pil_image)
                if result.success:
                    return result.text
                else:
                    logger.warning(f"OCR failed: {result.error}")

            return ""

        except Exception as e:
            logger.warning(f"OCR failed for page: {e}")
            return ""

    def _extract_from_docx(self, file_path: str) -> tuple[str, DocumentMetadata, int, bool]:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (text, metadata, page_count, ocr_used)
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

        doc = Document(file_path)
        metadata = DocumentMetadata()

        # Extract metadata from core properties
        if doc.core_properties:
            metadata.title = doc.core_properties.title
            metadata.author = doc.core_properties.author
            metadata.subject = doc.core_properties.subject
            if doc.core_properties.keywords:
                metadata.keywords = [k.strip() for k in doc.core_properties.keywords.split(",") if k.strip()]

        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        # Estimate page count (rough: ~500 words per page)
        full_text = "\n".join(text_parts)
        word_count = len(full_text.split())
        page_count = max(1, word_count // 500)

        return full_text, metadata, page_count, False

    def _extract_from_txt(self, file_path: str) -> tuple[str, DocumentMetadata, int, bool]:
        """Extract text from plain text file.

        Args:
            file_path: Path to text file

        Returns:
            Tuple of (text, metadata, page_count, ocr_used)
        """
        metadata = DocumentMetadata()
        metadata.title = Path(file_path).stem

        # Try different encodings
        for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            # Fallback with errors ignored
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

        # Estimate page count
        word_count = len(text.split())
        page_count = max(1, word_count // 500)

        return text, metadata, page_count, False

    def _extract_from_image(self, file_path: str, enable_ocr: bool) -> tuple[str, DocumentMetadata, int, bool]:
        """Extract text from image using Azure Document Intelligence.

        Args:
            file_path: Path to image file
            enable_ocr: Whether to perform OCR

        Returns:
            Tuple of (text, metadata, page_count, ocr_used)
        """
        metadata = DocumentMetadata()
        metadata.title = Path(file_path).stem

        if not enable_ocr:
            return "", metadata, 1, False

        # Use OCR provider manager
        ocr_manager = self._get_ocr_manager()
        if ocr_manager:
            result = ocr_manager.extract_from_image(file_path)
            if result.success:
                # Add OCR metadata
                if result.metadata:
                    metadata.custom_tags = metadata.custom_tags or []
                    provider = result.metadata.get("provider", "unknown")
                    metadata.custom_tags.append(f"ocr_provider:{provider}")

                    # Add table info if present
                    if result.tables:
                        metadata.custom_tags.append(f"tables:{len(result.tables)}")

                return result.text, metadata, 1, True
            else:
                logger.warning(f"OCR failed for {file_path}: {result.error}")

        return "", metadata, 1, False

    def chunk_text(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None,
    ) -> list[DocumentChunk]:
        """Split text into semantic chunks.

        Uses sentence boundaries and paragraph breaks to create
        natural chunk boundaries while respecting token limits.

        Args:
            text: Text to chunk
            chunk_size: Override default chunk size in tokens
            overlap: Override default overlap in tokens

        Returns:
            List of DocumentChunk objects
        """
        if not text or not text.strip():
            return []

        chunk_size = chunk_size or self.chunk_size_tokens
        overlap = overlap or self.chunk_overlap_tokens

        # Split into sentences (basic sentence tokenization)
        sentences = self._split_into_sentences(text)

        chunks = []
        current_chunk_sentences = []
        current_chunk_tokens = 0
        chunk_index = 0

        for sentence in sentences:
            sentence_tokens = self.count_tokens(sentence)

            # If single sentence exceeds chunk size, split it
            if sentence_tokens > chunk_size:
                # First, add current chunk if not empty
                if current_chunk_sentences:
                    chunk_text = " ".join(current_chunk_sentences)
                    chunks.append(DocumentChunk(
                        chunk_index=chunk_index,
                        chunk_text=chunk_text,
                        token_count=current_chunk_tokens,
                    ))
                    chunk_index += 1

                    # Keep overlap sentences
                    current_chunk_sentences = self._get_overlap_sentences(
                        current_chunk_sentences, overlap
                    )
                    current_chunk_tokens = self.count_tokens(" ".join(current_chunk_sentences))

                # Split the long sentence into smaller pieces
                long_sentence_chunks = self._split_long_sentence(sentence, chunk_size)
                for sub_chunk in long_sentence_chunks:
                    sub_tokens = self.count_tokens(sub_chunk)
                    chunks.append(DocumentChunk(
                        chunk_index=chunk_index,
                        chunk_text=sub_chunk,
                        token_count=sub_tokens,
                    ))
                    chunk_index += 1
                continue

            # Check if adding sentence would exceed chunk size
            if current_chunk_tokens + sentence_tokens > chunk_size and current_chunk_sentences:
                # Create chunk
                chunk_text = " ".join(current_chunk_sentences)
                chunks.append(DocumentChunk(
                    chunk_index=chunk_index,
                    chunk_text=chunk_text,
                    token_count=current_chunk_tokens,
                ))
                chunk_index += 1

                # Keep overlap sentences
                current_chunk_sentences = self._get_overlap_sentences(
                    current_chunk_sentences, overlap
                )
                current_chunk_tokens = self.count_tokens(" ".join(current_chunk_sentences))

            # Add sentence to current chunk
            current_chunk_sentences.append(sentence)
            current_chunk_tokens += sentence_tokens

        # Add final chunk
        if current_chunk_sentences:
            chunk_text = " ".join(current_chunk_sentences)
            chunks.append(DocumentChunk(
                chunk_index=chunk_index,
                chunk_text=chunk_text,
                token_count=self.count_tokens(chunk_text),
            ))

        return chunks

    def _split_into_sentences(self, text: str) -> list[str]:
        """Split text into sentences.

        Args:
            text: Text to split

        Returns:
            List of sentences
        """
        # Normalize whitespace
        text = re.sub(r"\s+", " ", text).strip()

        # Split on sentence boundaries
        # This regex handles common sentence endings including abbreviations
        sentence_pattern = r"(?<=[.!?])\s+(?=[A-Z])"
        sentences = re.split(sentence_pattern, text)

        # Filter empty sentences and strip whitespace
        return [s.strip() for s in sentences if s.strip()]

    def _get_overlap_sentences(self, sentences: list[str], overlap_tokens: int) -> list[str]:
        """Get sentences for overlap from end of chunk.

        Args:
            sentences: List of sentences in current chunk
            overlap_tokens: Target overlap in tokens

        Returns:
            List of sentences for overlap
        """
        if not sentences:
            return []

        overlap_sentences = []
        total_tokens = 0

        for sentence in reversed(sentences):
            sentence_tokens = self.count_tokens(sentence)
            if total_tokens + sentence_tokens <= overlap_tokens:
                overlap_sentences.insert(0, sentence)
                total_tokens += sentence_tokens
            else:
                break

        return overlap_sentences

    def _split_long_sentence(self, sentence: str, max_tokens: int) -> list[str]:
        """Split a long sentence into smaller chunks.

        Args:
            sentence: Long sentence to split
            max_tokens: Maximum tokens per chunk

        Returns:
            List of smaller text chunks
        """
        words = sentence.split()
        chunks = []
        current_words = []
        current_tokens = 0

        for word in words:
            word_tokens = self.count_tokens(word)
            if current_tokens + word_tokens > max_tokens and current_words:
                chunks.append(" ".join(current_words))
                current_words = []
                current_tokens = 0

            current_words.append(word)
            current_tokens += word_tokens

        if current_words:
            chunks.append(" ".join(current_words))

        return chunks

    def process_document(
        self,
        file_path: str,
        enable_ocr: bool = True,
        category: Optional[str] = None,
        custom_tags: Optional[list[str]] = None,
    ) -> RAGDocument:
        """Process a document for RAG ingestion.

        Args:
            file_path: Path to the document
            enable_ocr: Whether to use OCR for scanned documents
            category: Optional category for the document
            custom_tags: Optional list of custom tags

        Returns:
            RAGDocument with extracted text and chunks
        """
        file_path = str(Path(file_path).resolve())

        # Validate file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Get file info
        file_size = os.path.getsize(file_path)
        filename = os.path.basename(file_path)
        doc_type = self.get_document_type(file_path)

        if not doc_type:
            raise ValueError(f"Unsupported file type: {filename}")

        # Create document object
        doc = RAGDocument(
            filename=filename,
            file_type=doc_type,
            file_path=file_path,
            file_size_bytes=file_size,
            upload_status=UploadStatus.EXTRACTING,
        )

        try:
            # Extract text
            text, metadata, page_count, ocr_used = self.extract_text(file_path, enable_ocr)

            # Update document
            doc.page_count = page_count
            doc.ocr_required = ocr_used
            doc.metadata = metadata

            # Add custom metadata
            if category:
                doc.metadata.category = category
            if custom_tags:
                doc.metadata.custom_tags = custom_tags

            # Update status
            doc.upload_status = UploadStatus.CHUNKING

            # Chunk the text
            chunks = self.chunk_text(text)
            doc.chunks = chunks
            doc.chunk_count = len(chunks)

            # Mark chunking complete
            doc.upload_status = UploadStatus.EMBEDDING

        except Exception as e:
            doc.upload_status = UploadStatus.FAILED
            doc.error_message = str(e)
            logger.error(f"Failed to process document {filename}: {e}")

        return doc

    def compute_text_hash(self, text: str) -> str:
        """Compute SHA256 hash of text for caching.

        Args:
            text: Text to hash

        Returns:
            Hex-encoded SHA256 hash
        """
        return hashlib.sha256(text.encode("utf-8")).hexdigest()
